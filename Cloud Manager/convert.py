import os

def convert_word_to_pdf(input_file, output_file):
  """Converts a Word document (.docx) to PDF.

  Args:
    input_file: Path to the input Word document.
    output_file: Path to the output PDF file.
  """
  from docx import Document

  # Open the Word document
  doc = Document(input_file)

  # Save as PDF
  doc.save(output_file)

def convert_pdf_to_word(input_file, output_file):
  import PyPDF2

  # Open the PDF file
  with open(input_file, 'rb') as pdf_file:
    reader = PyPDF2.PdfReader(pdf_file)

  # Extract text from each page
  text = ""
  for page in reader.pages:
    text += page.extract_text()

  # Create a Word document and add the extracted text
  from docx import Document
  document = Document()
  document.add_paragraph(text)

  document.save(output_file)


input_file = "your_file.docx"  
output_file = "output.pdf"

from flask import Flask, request, jsonify
import os
import json

app = Flask(__name__)

@app.route('/convert', methods=['POST'])
def convert():
  if request.method == 'POST':
    # Get file and format data
    file = request.files['file']
    format = request.form['format']

    # Generate a unique filename
    filename, extension = os.path.splitext(file.filename)
    unique_filename = f"{filename}_{os.urandom(10).hex()}{extension}"

    # Save the uploaded file in a temporary directory
    upload_folder = 'temp'  # Replace with your desired temporary directory
    file_path = os.path.join(upload_folder, unique_filename)
    file.save(file_path)

    try:
      # Call the conversion function based on format
      if format == "pdf":
        convert_word_to_pdf(file_path, f"{unique_filename}.pdf")
        message = "File successfully converted to PDF."
      elif format == "docx":
        convert_pdf_to_word(file_path, f"{unique_filename}.docx")
        message = "File successfully converted to Word."
      else:
        message = "Invalid conversion format."
      
      # Optionally, delete the temporary file after successful conversion
      os.remove(file_path)
    except Exception as e:
      message = f"Error occurred during conversion: {str(e)}"

    # Prepare and return JSON response
    response = {'message': message}
    return jsonify(response)
convert_word_to_pdf(input_file, output_file)



